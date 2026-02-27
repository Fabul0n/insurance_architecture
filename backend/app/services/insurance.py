from datetime import datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from fastapi import HTTPException
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.dao.dao import InsuranceApplicationDAO, InsuranceContractDAO, InsurancePaymentDAO, InsuranceUserDAO
from app.schemas.insurance import (
    ApplicationCreateDB,
    ApplicationResponse,
    ContractCreateDB,
    PaymentCreateDB,
    PaymentRequest,
    PaymentResponse,
    RegisterRequest,
    UserCreateDB,
)
from app.services.auth import hash_password, verify_password
from db.models import InsuranceApplication, InsuranceContract, InsuranceUser


def cases_to_db(cases: list[str]) -> str:
    return "||".join(item.strip() for item in cases if item.strip())


def cases_from_db(cases: str) -> list[str]:
    return [item for item in cases.split("||") if item]


def get_public_content() -> dict:
    return {
        "title": "Страхование: зачем оно нужно",
        "paragraphs": [
            "Страхование помогает покрыть финансовые риски при непредвиденных событиях.",
            "Оно снижает нагрузку на личный бюджет при ущербе имуществу, здоровью или ответственности.",
            "В ряде случаев страхование обязательно по закону или условиям договора.",
        ],
        "note_for_guests": "Чтобы отправить заявку, сначала авторизуйтесь.",
    }


def get_personal_data_policy() -> dict:
    return {
        "title": "Политика обработки персональных данных",
        "updated_at": "2026-02-27",
        "sections": [
            "Мы обрабатываем персональные данные только для регистрации, оформления страховой заявки и исполнения договора.",
            "В состав обрабатываемых данных могут входить ФИО, паспортные данные, дата рождения, email и иные сведения из формы заявки.",
            "Данные не передаются третьим лицам, кроме случаев, предусмотренных законодательством и/или необходимых для исполнения договора.",
            "Пользователь подтверждает достоверность предоставленных данных и дает согласие на их обработку при установке соответствующего чекбокса.",
            "Пользователь вправе запросить уточнение, ограничение обработки или удаление своих данных через обращение в службу поддержки.",
        ],
    }


async def register_user(session: AsyncSession, body: RegisterRequest) -> InsuranceUser:
    if not any(ch.isdigit() for ch in body.password) or not any(ch.isalpha() for ch in body.password):
        raise HTTPException(status_code=400, detail="Пароль должен содержать буквы и цифры")

    existing = await InsuranceUserDAO.get_by_email(session, body.email, mute_not_found_exception=True)
    if existing:
        raise HTTPException(status_code=409, detail="Пользователь с таким email уже существует")

    return await InsuranceUserDAO.create_user(
        session,
        UserCreateDB(
            full_name=body.full_name,
            passport_data=body.passport_data,
            birth_date=body.birth_date,
            email=body.email,
            password_hash=hash_password(body.password),
        ),
    )


async def login_user(session: AsyncSession, email: str, password: str) -> InsuranceUser:
    user = await InsuranceUserDAO.get_by_email(session, email, mute_not_found_exception=True)
    if not user or not verify_password(password, user.password_hash):
        raise HTTPException(status_code=401, detail="Неверный email или пароль")
    return user


def application_to_response(application: InsuranceApplication) -> ApplicationResponse:
    return ApplicationResponse(
        id=application.id,
        status=application.status,
        full_name=application.full_name,
        passport_data=application.passport_data,
        birth_date=application.birth_date,
        email=application.email,
        workplace=application.workplace,
        insurance_object=application.insurance_object,
        insurance_period_months=application.insurance_period_months,
        insurance_cases=cases_from_db(application.insurance_cases),
        payout_amount=application.payout_amount,
        created_at=application.created_at,
    )


async def create_application(
    session: AsyncSession,
    user: InsuranceUser,
    *,
    full_name: str,
    passport_data: str,
    birth_date,
    email: str,
    workplace: str,
    insurance_object: str,
    insurance_period_months: int,
    insurance_cases: list[str],
    payout_amount: Decimal,
) -> InsuranceApplication:
    return await InsuranceApplicationDAO.create_application(
        session,
        ApplicationCreateDB(
            user_id=user.id,
            full_name=full_name,
            passport_data=passport_data,
            birth_date=birth_date,
            email=email,
            workplace=workplace,
            insurance_object=insurance_object,
            insurance_period_months=insurance_period_months,
            insurance_cases=cases_to_db(insurance_cases),
            payout_amount=payout_amount,
            status="awaiting_payment",
        ),
    )


async def get_last_application(session: AsyncSession, user_id: int) -> InsuranceApplication | None:
    return await InsuranceApplicationDAO.get_last_for_user(session, user_id)


async def list_applications(session: AsyncSession, user_id: int) -> list[InsuranceApplication]:
    return await InsuranceApplicationDAO.list_for_user(session, user_id)


async def make_payment(
    session: AsyncSession,
    user: InsuranceUser,
    body: PaymentRequest,
) -> PaymentResponse:
    application = await InsuranceApplicationDAO.get_for_user(session, body.application_id, user.id)
    if not application:
        raise HTTPException(status_code=404, detail="Заявка не найдена")
    if application.status == "paid":
        raise HTTPException(status_code=400, detail="Заявка уже оплачена")

    clean_card = (body.card_number or "").replace(" ", "")
    if body.payment_method == "card":
        if len(clean_card) < 16:
            raise HTTPException(status_code=400, detail="Введите корректный номер карты")
        card_last4 = clean_card[-4:]
    else:
        card_last4 = None

    failed = body.payment_method == "card" and clean_card.endswith("0000")
    payment_status = "failed" if failed else "success"

    payment = await InsurancePaymentDAO.create_payment(
        session,
        PaymentCreateDB(
            application_id=application.id,
            payment_method=body.payment_method,
            card_last4=card_last4,
            amount=Decimal(application.payout_amount),
            status=payment_status,
            fail_reason="Транзакция отклонена банком" if failed else None,
        ),
    )

    if failed:
        application.status = "payment_failed"
        return PaymentResponse(
            payment_id=payment.id,
            application_id=application.id,
            status="failed",
            message="Оплата неуспешна. Проверьте данные карты и попробуйте снова.",
        )

    application.status = "paid"
    await InsuranceContractDAO.create_contract(
        session,
        ContractCreateDB(
            user_id=user.id,
            application_id=application.id,
            contract_number=f"INS-{datetime.utcnow():%Y%m%d}-{str(uuid4())[:8].upper()}",
        ),
    )
    return PaymentResponse(
        payment_id=payment.id,
        application_id=application.id,
        status="success",
        message="Оплата прошла успешно. Договор доступен в личном кабинете.",
    )


async def list_contracts(session: AsyncSession, user_id: int) -> list[InsuranceContract]:
    return await InsuranceContractDAO.list_for_user(session, user_id)


async def get_contract_for_user(
    session: AsyncSession, contract_id: int, user_id: int
) -> tuple[InsuranceContract, InsuranceApplication]:
    contract = await InsuranceContractDAO.get_for_user(session, contract_id, user_id)
    if not contract:
        raise HTTPException(status_code=404, detail="Договор не найден")
    application = await InsuranceApplicationDAO.get_for_user(session, contract.application_id, user_id)
    if not application:
        raise HTTPException(status_code=404, detail="Связанная заявка не найдена")
    return contract, application


def _replace_tokens_in_runs(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        for token, value in replacements.items():
            if token in paragraph_text:
                paragraph_text = paragraph_text.replace(token, value)
        if paragraph_text != paragraph.text:
            paragraph.text = paragraph_text

        for run in paragraph.runs:
            for token, value in replacements.items():
                if token in run.text:
                    run.text = run.text.replace(token, value)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text = paragraph.text
                    for token, value in replacements.items():
                        if token in paragraph_text:
                            paragraph_text = paragraph_text.replace(token, value)
                    if paragraph_text != paragraph.text:
                        paragraph.text = paragraph_text

                    for run in paragraph.runs:
                        for token, value in replacements.items():
                            if token in run.text:
                                run.text = run.text.replace(token, value)


def build_contract_docx_from_template(
    contract: InsuranceContract,
    application: InsuranceApplication,
) -> bytes:
    settings = get_settings()
    template_path = Path(settings.CONTRACT_TEMPLATE_PATH)
    if not template_path.exists():
        raise HTTPException(
            status_code=500,
            detail=f"Шаблон договора не найден: {template_path}",
        )

    document = Document(str(template_path))
    replacements = {
        "{{CONTRACT_NUMBER}}": contract.contract_number,
        "{{CONTRACT_DATE}}": contract.created_at.strftime("%d.%m.%Y"),
        "{{POLICYHOLDER_FULL_NAME}}": application.full_name,
        "{{POLICYHOLDER_PASSPORT}}": application.passport_data,
        "{{POLICYHOLDER_BIRTH_DATE}}": application.birth_date.strftime("%d.%m.%Y"),
        "{{POLICYHOLDER_EMAIL}}": application.email,
        "{{WORKPLACE}}": application.workplace,
        "{{INSURANCE_OBJECT}}": application.insurance_object,
        "{{INSURANCE_PERIOD_MONTHS}}": str(application.insurance_period_months),
        "{{INSURANCE_CASES}}": ", ".join(cases_from_db(application.insurance_cases)),
        "{{PAYOUT_AMOUNT}}": str(application.payout_amount),
        "{{INSURER_FULL_NAME}}": settings.INSURER_FULL_NAME,
        "<НОМЕР_ДОГОВОРА>": contract.contract_number,
        "<ДАТА_ДОГОВОРА>": contract.created_at.strftime("%d.%m.%Y"),
        "<ФИО_СТРАХОВАТЕЛЯ>": application.full_name,
        "<ПАСПОРТ_СТРАХОВАТЕЛЯ>": application.passport_data,
        "<ДАТА_РОЖДЕНИЯ_СТРАХОВАТЕЛЯ>": application.birth_date.strftime("%d.%m.%Y"),
        "<EMAIL_СТРАХОВАТЕЛЯ>": application.email,
        "<МЕСТО_РАБОТЫ>": application.workplace,
        "<ОБЪЕКТ_СТРАХОВАНИЯ>": application.insurance_object,
        "<СРОК_СТРАХОВАНИЯ_МЕС>": str(application.insurance_period_months),
        "<СТРАХОВЫЕ_СЛУЧАИ>": ", ".join(cases_from_db(application.insurance_cases)),
        "<СУММА_ВЫПЛАТЫ>": str(application.payout_amount),
        "<ФИО_СТРАХОВЩИКА>": settings.INSURER_FULL_NAME,
    }
    _replace_tokens_in_runs(document, replacements)

    out = BytesIO()
    document.save(out)
    out.seek(0)
    return out.read()


def _extract_docx_lines(docx_bytes: bytes) -> list[str]:
    document = Document(BytesIO(docx_bytes))
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            row_parts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_parts:
                lines.append(" | ".join(row_parts))

    return lines


def build_contract_pdf_from_template(
    contract: InsuranceContract,
    application: InsuranceApplication,
) -> bytes:
    # PDF is generated from the already-filled DOCX template.
    docx_bytes = build_contract_docx_from_template(contract, application)
    lines = _extract_docx_lines(docx_bytes)

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    body_style = styles["BodyText"]
    font_name = _ensure_cyrillic_font()
    body_style.fontName = font_name
    body_style.fontSize = 10
    body_style.leading = 13

    story = [Paragraph(line.replace("\n", "<br/>"), body_style) for line in lines]
    spaced_story: list = []
    for item in story:
        spaced_story.append(item)
        spaced_story.append(Spacer(1, 6))

    doc.build(spaced_story)
    buffer.seek(0)
    return buffer.read()


def _ensure_cyrillic_font() -> str:
    preferred_name = "DejaVuSans"
    if preferred_name in pdfmetrics.getRegisteredFontNames():
        return preferred_name

    candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/local/share/fonts/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont(preferred_name, str(candidate)))
            return preferred_name

    return "Helvetica"
